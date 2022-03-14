import os
import pandas as pd
from docx2pdf import convert
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pytesseract as pt
from pdf2image import convert_from_path
import re

pt.pytesseract.tesseract_cmd = 'C:/Program Files/Tesseract-OCR/tesseract.exe'

class DataDigitization:

    def __init__(self, root):
        self.root = root
        self.report_names_df = pd.read_excel(os.path.join(self.root, 'source/Report_types_17.xlsx'))
        self.master_list = pd.read_excel(os.path.join(self.root, 'source/2022_03_07_patient_master_list_dummy.xlsx'))
        self.coded_data = os.path.join(self.root, 'output/coded_data')
        self.qr_code_folder_path = os.path.join(self.root, 'sample_from_HR/549_16')
        self.tmp = os.path.join(self.root, 'tmp')
        self.scanned_patient_files = os.path.join(self.root, 'scanned_patient_files\original_pdf')
        self.scanned_file_images = os.path.join(self.root, 'scanned_patient_files\spplitted_pdf_to_img')

    def add_qr_code(self):
        qr_codes = os.listdir(self.qr_code_folder_path)
        for qr_code in qr_codes:
            doc = Document()
            qr_code_path = os.path.join(self.qr_code_folder_path, qr_code)
            doc.add_picture(qr_code_path)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            qr_code = str(qr_code)
            qr_code = qr_code.lower()
            report_name = re.sub('[^a-z_.]', '', str(qr_code))
            report_name = re.sub('.png', '', report_name)
            report_name = report_name[2:]
            text = doc.add_paragraph()
            report_type_name = text.add_run(report_name)
            report_type_name.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            report_type_name.bold = True
            report_type_name.font.size = Pt(28)
            report_type_name.font.name = 'Arial Black'
            for i in range(len(self.master_list)):
                file_no = 'File_number: ' + str(self.master_list['file_number'][i])
                mr_no = 'MR_number: ' + str(self.master_list['mr_number'][i])
                patient_name = 'Patient_name: ' + str(self.master_list['patient_name'][i])
                dob = 'Date_of_Birth: ' + str(self.master_list['date_of_birth'][i])
                dir = self.master_list['file_number'][i]
                dir = re.sub('/', '_', str(dir))
                dir_path = os.path.join(self.coded_data, dir)
                if not os.path.isdir(dir_path):
                    os.mkdir(dir_path)
                blank_para = doc.add_paragraph()
                run = blank_para.add_run()
                run.add_break()
                id = doc.add_paragraph().add_run(file_no)
                id.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                id = doc.add_paragraph().add_run(mr_no)
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                id = doc.add_paragraph().add_run(patient_name)
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                id = doc.add_paragraph().add_run(dob)
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                report_type = re.sub('.png', '.docx', str(qr_code))
                report_type = report_type.lower()
                doc_name = report_type + '.docx'
                doc.save(os.path.join(dir_path, doc_name))
                return dir_path

    @staticmethod
    def create_dummy_pdf_for_reports(report_type, dir_path):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', size=28)
        text_letter = re.sub('[^a-z_]', '', str(report_type))
        for i in range(5):
            text = text_letter[2:] + '_' + str(i)
            pdf.cell(200, 10, txt=text, ln=1, align='C')
            pdf.add_page()
            report_type = re.sub('[^a-z_]', '', str(report_type))
        pdf_name = report_type[2:] + '.pdf'
        pdf.output(os.path.join(dir_path, pdf_name))

    ## this function creates a 5 different pages of reports

    @staticmethod
    def create_dummy_images_for_reports(report_type, dir_path):
        for i in range(5):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('Arial', size=28)
            text_letter = re.sub('[^a-z_]', '', str(report_type))
            print(text_letter)
            text = text_letter[1:] + '_' + str(i)
            pdf.cell(200, 10, txt=text, ln=1, align='C')
            report_type = re.sub('[^a-z_]', '', str(report_type))
            pdf_name = report_type[1:] + '_' + str(i) + '.pdf'
            pdf.output(os.path.join(dir_path, pdf_name))

    ## this function creates a separate directory for each report type and then add the dummy data in it

    def add_report_data(self, qr_label):
        reports = os.listdir(qr_label)
        for report in reports:
            report_name = re.sub('.docx', '', str(report))
            each_report_dir = os.path.join(self.tmp, report_name)
            if not os.path.isdir(each_report_dir):
                os.mkdir(each_report_dir)
            doc_path = os.path.join(self.tmp, report_name)
            pdf_report_name = re.sub('[^0-9]', '', str(report_name))
            pdf_report_name = pdf_report_name + '_code.pdf'
            pdf_path = os.path.join(each_report_dir, pdf_report_name)
            convert(doc_path, pdf_path)
            self.create_dummy_images_for_reports(report_name, each_report_dir)

    ## this function creates a dummy data into same folder
    def add_dummy_reports_(self):
        reports = os.listdir(self.coded_data)
        for report in reports:
            report_name = re.sub('.docx', '', str(report))
            doc_path = os.path.join(self.coded_data, report)
            # pdf_report_name = re.sub('[^0-9_]', '', str(report))
            pdf_report_name = report_name + '_code.pdf'
            pdf_path = os.path.join(self.tmp, pdf_report_name)
            convert(doc_path, pdf_path)
            self.create_dummy_pdf_for_reports(report_name, self.tmp)

    def split_pdf_by_images(self):
        scanned_files = os.listdir(self.scanned_patient_files)
        for scanned_file in scanned_files:
            print(scanned_file)
            scanned_file = scanned_file.lower()
            file_num = re.sub('[^0-9_]', '', scanned_file)
            splitted_file_path = os.path.join(self.scanned_file_images, file_num)
            if not os.path.isdir(splitted_file_path):
                os.mkdir(splitted_file_path)
            if scanned_file.endswith('.pdf'):
                pages = convert_from_path(os.path.join(self.scanned_file_images, scanned_file), 500,
                                          poppler_path='C:/Program Files/poppler-0.68.0/bin')
                i = 0
                for index, page in enumerate(pages):
                    if i == index:
                        file_name = scanned_file.lower()
                        file_name = re.sub('.pdf', '', file_name)
                        page_no = i + 1
                        out_jpg = file_name + '_' + str(page_no) + '.jpg'
                        page.save(os.path.join(splitted_file_path, out_jpg), 'JPEG')
                    i += 1
                print("file number: ", file_num + " split")

        # def extract_report_data_from_file(self):
        # scanned data either as images or as pdf
        # excel of categorised data - dk to send now
