import os
import re
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx2pdf import convert
from fpdf import FPDF

qr_code_folder_path = 'D:/Shweta/data_digitization/sample_from_HR/549_16'
report_names_df = pd.read_excel('D:/Shweta/data_digitization/reference_docs/Report_types_17.xlsx',
                                sheet_name='report_types')
master_list = pd.read_excel('D:/Shweta/data_digitization/reference_docs/Report_types_17.xlsx',
                            sheet_name='master_list')
destination_path = 'D:/Shweta/data_digitization/sample_from_HR/report_types_added_qr_code'

def add_qr_code_in_word_doc(report_type_df_path, qr_code_path, destination_path, master_list):
    for index, report_type in enumerate(report_type_df_path['report_types']):
        report_type_no = index + 1
        doc = Document()
        doc.add_picture(qr_code_path)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        type = report_type
        text = doc.add_paragraph()
        report_type_name = text.add_run(str(report_type_no) + '. ' + type)
        report_type_name.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        report_type_name.bold = True
        report_type_name.font.size = Pt(28)
        report_type_name.font.name = 'Arial Black'
        # doc.save(os.path.join(destination_path, report_type + '.docx'))
        for i in range(len(master_list)):
            file_no = 'File_number: ' + str(master_list['file_number'][i])
            mr_no = 'MR_number: ' + str(master_list['mr_number'][i])
            patient_name = 'Patient_name: ' + str(master_list['patient_name'][i])
            dob = 'Date_of_Birth: ' + str(master_list['date_of_birth'][i])
            dir = master_list['file_number'][i]
            dir = re.sub('/', '_', str(dir))
            dir_path = os.path.join(destination_path, dir)
            if not os.path.isdir(dir_path):
                os.mkdir(dir_path)
            blank_para = doc.add_paragraph()
            run = blank_para.add_run()
            run.add_break()
            id = doc.add_paragraph().add_run(file_no)
            id.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(mr_no)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(patient_name)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(dob)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            report_type = re.sub(' ', '_', report_type)
            report_type = report_type.lower()
            doc_name = str(report_type_no) + '_' + report_type + '.docx'
            doc.save(os.path.join(dir_path, doc_name))

# add_qr_code_in_word_doc(report_names_df, dummy_img, destination_path, master_list)

##

def create_dummy_images_for_reports(report_type, dir_path):
    for i in range(5):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', size=28)
        text_letter = re.sub('[^a-z_]', '', str(report_type))
        print(text_letter)
        text = text_letter[1:] + '_' + str(i)
        pdf.cell(200, 10, txt=text, ln=1, align='C')
        report_type = re.sub('[^a-z_]', '', str(report_type))
        pdf_name = report_type[1:] + '_' + str(i) + '.pdf'
        pdf.output(os.path.join(dir_path, pdf_name))

def add_dummy_reports(source_path, destination_path):
    reports = os.listdir(source_path)
    for report in reports:
        print(report)
        dir_name = re.sub('.docx', '', str(report))
        print(dir_name)
        each_report_dir = os.path.join(destination_path, dir_name)
        os.mkdir(each_report_dir)
        doc_path = os.path.join(source_path, report)
        print(doc_path)
        pdf_report_name = re.sub('[^0-9]', '', str(report))
        pdf_report_name = pdf_report_name + '_code.pdf'
        pdf_path = os.path.join(each_report_dir, pdf_report_name)
        convert(doc_path, pdf_path)
        create_dummy_images_for_reports(dir_name, each_report_dir)

add_dummy_reports('D:/Shweta/data_digitization/sample_output/dummy_report_types/12_34',
                  'D:/Shweta/data_digitization/sample_output/dummy_report_types/dummy_files')

## adding original_qr_code

def add_qr_code(qr_folder_path, master_list, destination_path):
    qr_codes = os.listdir(qr_folder_path)
    for qr_code in qr_codes:
        doc = Document()
        qr_code_path = os.path.join(qr_folder_path, qr_code)
        doc.add_picture(qr_code_path)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        qr_code = str(qr_code)
        qr_code = qr_code.lower()
        report_name = re.sub('[^a-z_.]', '', str(qr_code))
        report_name = re.sub('.png', '', report_name)
        report_name = report_name[2:]
        text = doc.add_paragraph()
        report_type_name = text.add_run(report_name)
        report_type_name.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        report_type_name.bold = True
        report_type_name.font.size = Pt(28)
        report_type_name.font.name = 'Arial Black'
        for i in range(len(master_list)):
            file_no = 'File_number: ' + str(master_list['file_number'][i])
            mr_no = 'MR_number: ' + str(master_list['mr_number'][i])
            patient_name = 'Patient_name: ' + str(master_list['patient_name'][i])
            dob = 'Date_of_Birth: ' + str(master_list['date_of_birth'][i])
            dir = master_list['file_number'][i]
            dir = re.sub('/', '_', str(dir))
            dir_path = os.path.join(destination_path, dir)
            if not os.path.isdir(dir_path):
                os.mkdir(dir_path)
            blank_para = doc.add_paragraph()
            run = blank_para.add_run()
            run.add_break()
            id = doc.add_paragraph().add_run(file_no)
            id.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(mr_no)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(patient_name)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(dob)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            report_type = re.sub('.png', '.docx', str(qr_code))
            # report_type = report_type.lower()
            # doc_name = report_type + '.docx'
            doc.save(os.path.join(dir_path, report_type))

##

def create_dummy_pdf_for_reports(report_type, dir_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', size=28)
    text_letter = re.sub('[^a-z_]', '', str(report_type))
    print(text_letter)
    for i in range(5):
        text = text_letter[2:] + '_' + str(i)
        pdf.cell(200, 10, txt=text, ln=1, align='C')
        pdf.add_page()
        report_type = re.sub('[^a-z_]', '', str(report_type))
    pdf_name = report_type[2:] + '.pdf'
    pdf.output(os.path.join(dir_path, pdf_name))

def add_dummy_reports_(source_path, destination_path):
    reports = os.listdir(source_path)
    for report in reports:
        print(report)
        dir_name = re.sub('.docx', '', str(report))
        print(dir_name)
        doc_path = os.path.join(source_path, report)
        print(doc_path)
        # pdf_report_name = re.sub('[^0-9_]', '', str(report))
        pdf_report_name = dir_name + '_code.pdf'
        pdf_path = os.path.join(destination_path, pdf_report_name)
        convert(doc_path, pdf_path)
        create_dummy_pdf_for_reports(dir_name, destination_path)


# add_dummy_reports_('D:/Shweta/data_digitization/sample_from_HR/report_types_added_qr_code/549_16',
#                    'D:/Shweta/data_digitization/sample_from_HR/report_types_added_qr_code/added_dummy_reports')

