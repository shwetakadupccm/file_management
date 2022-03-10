import os
import re
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx2pdf import convert
from fpdf import FPDF
import pyqrcode
import datetime

qr_code_folder_path = 'D:/Shweta/data_digitization/sample_from_HR/549_16'
report_names_df = pd.read_excel('D:/Shweta/data_digitization/reference_docs/Report_types_17.xlsx',
                                sheet_name='report_types')
master_list = pd.read_excel('D:/Shweta/data_digitization/reference_docs/2022_03_07_patient_master_list_dummy.xlsx')
destination_path = 'D:/Shweta/data_digitization/sample_output/2022_03_10/added_id_info/added_dummy_reports'
coded_data = 'D:/Shweta/data_digitization/sample_output/2022_03_10/added_id_info/549_16'

def add_qr_code(dummy_qr_folder_path, master_list, destination_path):
    qr_codes = os.listdir(dummy_qr_folder_path)
    for qr_code in qr_codes:
        doc = Document()
        qr_code_path = os.path.join(dummy_qr_folder_path, qr_code)
        doc.add_picture(qr_code_path)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        qr_code = str(qr_code)
        qr_code = qr_code.lower()
        report_name = re.sub('[^a-z_.]', '', str(qr_code))
        report_name = re.sub('.png', '', report_name)
        report_name = report_name[4:]
        text = doc.add_paragraph()
        report_type_name = text.add_run(report_name)
        report_type_name.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        report_type_name.bold = True
        report_type_name.font.size = Pt(28)
        report_type_name.font.name = 'Arial Black'
        for i in range(len(master_list)):
            file_no = 'File_number: ' + str(master_list['file_number'][i])
            mr_no = 'MR_number: ' + str(master_list['mr_number'][i])
            patient_name = 'Patient_name: ' + str(master_list['patient_name'][i])
            dob = 'Date_of_Birth: ' + str(master_list['date_of_birth'][i])
            dir = master_list['file_number'][i]
            dir = re.sub('/', '_', str(dir))
            dir_path = os.path.join(destination_path, dir)
            if not os.path.isdir(dir_path):
                os.mkdir(dir_path)
            blank_para = doc.add_paragraph()
            run = blank_para.add_run()
            run.add_break()
            id = doc.add_paragraph().add_run(file_no)
            id.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(mr_no)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(patient_name)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(dob)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            report_type = re.sub('.png', '.docx', str(qr_code))
            # report_type = report_type.lower()
            # doc_name = report_type + '.docx'
            doc.save(os.path.join(dir_path, report_type))

add_qr_code('D:/Shweta/data_digitization/qr_code_creation/qr_codes_by_subfolder',
            master_list, 'D:/Shweta/data_digitization/sample_output/2022_03_10/added_id_info/coded_data_for_subfolders')

def create_dummy_pdf_for_reports(report_type, dir_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', size=28)
    text_letter = re.sub('[^a-z_]', '', str(report_type))
    print(text_letter)
    for i in range(5):
        report_no = i+1
        text = text_letter[2:] + '_' + str(report_no)
        pdf.cell(200, 10, txt=text, ln=1, align='C')
        pdf.add_page()
        report_type = re.sub('[^a-z_]', '', str(report_type))
    pdf_name = report_type[2:] + '.pdf'
    pdf.output(os.path.join(dir_path, pdf_name))

def add_dummy_reports(source_path, destination_path):
    reports = os.listdir(source_path)
    for report in reports:
        print(report)
        dir_name = re.sub('.docx', '', str(report))
        print(dir_name)
        each_report_dir = os.path.join(destination_path, dir_name)
        os.mkdir(each_report_dir)
        doc_path = os.path.join(source_path, report)
        print(doc_path)
        pdf_report_name = re.sub('[^0-9_]', '', str(report))
        print(pdf_report_name)
        pdf_report_name = pdf_report_name[:-1] + '_code.pdf'
        pdf_path = os.path.join(each_report_dir, pdf_report_name)
        convert(doc_path, pdf_path)
        report_type = re.sub('[^a-z_]', '', str(dir_name))
        if report_type[2:] == 'patient_information':
            sub_dir_pccm_form = os.path.join(each_report_dir, 'PCCM_form')
            os.mkdir(sub_dir_pccm_form)
            create_dummy_pdf_for_reports('__pccm_form', sub_dir_pccm_form)
            sub_dir_id_proofs = os.path.join(each_report_dir, 'ID_proofs')
            os.mkdir(sub_dir_id_proofs)
            create_dummy_pdf_for_reports('__id_proof', sub_dir_id_proofs)
        elif report_type[2:] == 'radiology':
            sub_dir_screening = os.path.join(each_report_dir, 'screening')
            os.mkdir(sub_dir_screening)
            create_dummy_pdf_for_reports('__screening', sub_dir_screening)
            sub_dir_diagnosis = os.path.join(each_report_dir, 'diagnosis')
            os.mkdir(sub_dir_diagnosis)
            create_dummy_pdf_for_reports('__diagnosis', sub_dir_diagnosis)
            sub_dir_nact = os.path.join(each_report_dir, 'nact')
            os.mkdir(sub_dir_nact)
            sub_dir_nact_pre = os.path.join(sub_dir_nact, 'pre')
            os.mkdir(sub_dir_nact_pre)
            create_dummy_pdf_for_reports('__pre_nact', sub_dir_nact_pre)
            sub_dir_nact_post = os.path.join(sub_dir_nact, 'post')
            os.mkdir(sub_dir_nact_post)
            create_dummy_pdf_for_reports('__post_nact', sub_dir_nact_post)
        else:
            create_dummy_pdf_for_reports(dir_name, each_report_dir)

add_dummy_reports(coded_data, destination_path)

