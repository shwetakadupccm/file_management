import math
import numpy as np
import pandas as pd
# from PyPDF2 import PdfFileReader
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pytesseract as pt
# from PIL import Image
from pdf2image import convert_from_path
import re
import shutil
pt.pytesseract.tesseract_cmd = 'C:/Program Files/Tesseract-OCR/tesseract.exe'

# file_path = 'D:/Shweta/data_digitization/scanned_patient_files/15_19.PDF'
root = 'D:/Shweta/data_digitization'
report_names_df = pd.read_excel(os.path.join(root, 'reference_docs/Report_types_17.xlsx'))
master_list = pd.read_excel(os.path.join(root, 'reference_docs/2022_03_07_patient_master_list_dummy.xlsx'))
coded_data = os.path.join(root, 'sample_output/2022_03_07_coded_data')
qr_code_path = os.path.join(root, 'sample_from_HR/549_16')
tmp_folder = os.path.join(root, 'tmp')
scanned_files = os.path.join(root, 'scanned_patient_files\original_pdf')
splitted_scanned_files = os.path.join(root, 'scanned_patient_files/spplitted_pdf_to_img')
scanned_file_log_df = pd.read_excel(os.path.join(root, 'reference_docs/2022_02_09_Data_digitzation_scanned_files_dk.xlsx'))
col_heads = pd.read_excel(os.path.join(root, 'reference_docs/2022_02_09_Data_digitzation_scanned_files_dk.xlsx'),
                          sheet_name='col_heads')
## adding qr code

def add_qr_code_id_data(qr_code_folder_path, master_list, coded_data):
    qr_codes = os.listdir(qr_code_folder_path)
    for qr_code in qr_codes:
        doc = Document()
        qr_code_path = os.path.join(qr_code_folder_path, qr_code)
        doc.add_picture(qr_code_path)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        qr_code = str(qr_code)
        qr_code = qr_code.lower()
        report_name = re.sub('[^a-z_.]', '', str(qr_code))
        report_name = re.sub('.png', '', report_name)
        report_name = report_name[2:]
        text = doc.add_paragraph()
        report_type_name = text.add_run(report_name)
        report_type_name.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        report_type_name.bold = True
        report_type_name.font.size = Pt(28)
        report_type_name.font.name = 'Arial Black'
        for i in range(len(master_list)):
            file_no = 'File_number: ' + str(master_list['file_number'][i])
            mr_no = 'MR_number: ' + str(master_list['mr_number'][i])
            patient_name = 'Patient_name: ' + str(master_list['patient_name'][i])
            dob = 'Date_of_Birth: ' + str(master_list['date_of_birth'][i])
            dir = master_list['file_number'][i]
            dir = re.sub('/', '_', str(dir))
            dir_path = os.path.join(coded_data, dir)
            if not os.path.isdir(dir_path):
                os.mkdir(dir_path)
            blank_para = doc.add_paragraph()
            run = blank_para.add_run()
            run.add_break()
            id = doc.add_paragraph().add_run(file_no)
            id.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(mr_no)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(patient_name)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            id = doc.add_paragraph().add_run(dob)
            id.font.size = Pt(20)
            id.font.name = 'Arial Black'
            report_type = re.sub('.png', '.docx', str(qr_code))
            report_type = report_type.lower()
            doc_name = report_type
            doc.save(os.path.join(dir_path, doc_name))

## converting pdf to jpg image

def split_pdf_by_images(scanned_files_path, splitted_scan_files):
    scanned_files = os.listdir(scanned_files_path)
    for scanned_file in scanned_files:
        print(scanned_file)
        scanned_file = scanned_file.lower()
        file_num = re.sub('[^0-9_]', '', scanned_file)
        splitted_file_path = os.path.join(splitted_scan_files, file_num)
        if not os.path.isdir(splitted_file_path):
            os.mkdir(splitted_file_path)
        if scanned_file.endswith('.pdf'):
            pages = convert_from_path(os.path.join(scanned_files_path, scanned_file), 500,
                                      poppler_path = 'C:/Program Files/poppler-0.68.0/bin')
            i = 0
            for index, page in enumerate(pages):
                if i == index:
                    file_name = scanned_file.lower()
                    file_name = re.sub('.pdf', '', file_name)
                    page_no = i + 1
                    out_jpg = file_name + '_' + str(page_no) + '.jpg'
                    page.save(os.path.join(splitted_file_path, out_jpg), 'JPEG')
                i += 1
            print("file number: ", file_num + " split")

split_pdf_by_images('D:/Shweta/data_digitization/scanned_patient_files/original_pdf',
                    'D:/Shweta/data_digitization/scanned_patient_files/spplitted_pdf_to_img')

##
def get_image_no(file_number, file_images_lst):
    file_images_no_lst = []
    for file_image in file_images_lst:
        file_image_no = re.sub(file_number, '', str(file_image))
        file_image_no = re.sub('.jpg', '', file_image_no)
        file_image_no = re.sub('_', '', file_image_no)
        file_image_no = file_image_no.strip()
        file_images_no_lst.append(file_image_no)
    return file_images_no_lst

def split_report_page_no(report_page_no):
    if isinstance(report_page_no, float):
        page_no_lst = []
        if not math.isnan(report_page_no):
            integer = int(report_page_no)
            page_no_lst.append(str(integer))
        return page_no_lst
    elif isinstance(report_page_no, int):
        page_no_lst = []
        page_no_lst.append(str(report_page_no))
        return page_no_lst
    elif ';' in report_page_no:
        page_no_lst = []
        report_page_no_splitted = report_page_no.split(';')
        for page_no in report_page_no_splitted:
            if '|' in page_no:
                partitions = page_no.partition('|')
                start = int(partitions[0])
                end = int(partitions[2]) + 1
                page_nos = np.arange(start, end)
                page_nos_lst = page_nos.tolist()
                for no in page_nos_lst:
                    page_no_lst.append(str(no))
            else:
                page_no_lst.append(str(page_no))
        return page_no_lst
    elif '|' in report_page_no:
        page_no_lst = []
        partitions = report_page_no.partition('|')
        start = int(partitions[0])
        end = int(partitions[2]) + 1
        page_nos = np.arange(start, end)
        page_nos_lst = page_nos.tolist()
        for no in page_nos_lst:
            page_no_lst.append(str(no))
        return page_no_lst
    elif type(report_page_no) in (float, int):
        page_no_lst = []
        report_page_no = int(report_page_no)
        page_no_lst.append(str(report_page_no))
        return page_no_lst
    else:
        page_no_lst = []
        page_no_lst.append(str(report_page_no))
        return page_no_lst

def categorize_scanned_file_by_report_types(splitted_scanned_files_path, file_categorization_df, file_categorization_df_col_heads):
    splitted_scanned_files = os.listdir(splitted_scanned_files_path)
    report_types = file_categorization_df_col_heads['report_types']
    for index, file_number in enumerate(file_categorization_df['File Number']):
        if file_number in splitted_scanned_files:
            img_lst = os.listdir(os.path.join(splitted_scanned_files_path, file_number))
            img_no_lst = get_image_no(file_number, img_lst)
            for report_type in report_types:
                print(report_type)
                report_page_no = file_categorization_df.iloc[index][report_type]
                print(report_page_no)
                report_page_no_splitted = split_report_page_no(report_page_no)
                file_no_dir = os.path.join(splitted_scanned_files_path, file_number)
                report_dir = os.path.join(file_no_dir, report_type)
                if not os.path.isdir(report_dir):
                    os.mkdir(report_dir)
                for page_no in report_page_no_splitted:
                    if page_no in img_no_lst:
                        report_page_name = str(file_number) + '_' + str(page_no) + '.jpg'
                        print(report_page_name)
                        all_pages_dir = os.path.join(splitted_scanned_files_path, str(file_number))
                        source_path = os.path.join(all_pages_dir, report_page_name)
                        destination_path = os.path.join(report_dir, report_page_name)
                        shutil.copy(source_path, destination_path)
                        print('image_moved')

categorize_scanned_file_by_report_types(splitted_scanned_files, scanned_file_log_df, col_heads)

