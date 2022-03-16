import os
import re
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx2pdf import convert
from fpdf import FPDF

class DataDigitization:
    def __init__(self, report_type_df, qr_code_folder_path, master_list, qr_code_id_folder_path, destination_path):
        self.report_type_df = report_type_df
        self.qr_code_folder_path = qr_code_folder_path
        self.master_list = master_list
        self.qr_code_id_folder_path = qr_code_id_folder_path
        self.destination_path = destination_path

    def add_qr_code(self):
        qr_codes = os.listdir(self.qr_code_folder_path)
        for qr_code in qr_codes:
            doc = Document()
            qr_code_path = os.path.join(self.qr_code_folder_path, qr_code)
            doc.add_picture(qr_code_path)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            qr_code = str(qr_code)
            qr_code = qr_code.lower()
            report_name = re.sub('[^a-z_.]', '', str(qr_code))
            report_name = re.sub('.png', '', report_name)
            report_name = report_name[2:]
            text = doc.add_paragraph()
            report_type_name = text.add_run(report_name)
            report_type_name.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            report_type_name.bold = True
            report_type_name.font.size = Pt(28)
            report_type_name.font.name = 'Arial Black'
            for i in range(len(self.master_list)):
                file_no = 'File_number: ' + str(self.master_list['file_number'][i])
                mr_no = 'MR_number: ' + str(self.master_list['mr_number'][i])
                patient_name = 'Patient_name: ' + str(self.master_list['patient_name'][i])
                dob = 'Date_of_Birth: ' + str(self.master_list['date_of_birth'][i])
                dir = self.master_list['file_number'][i]
                dir = re.sub('/', '_', str(dir))
                dir_path = os.path.join(self.qr_code_id_folder_path, dir)
                if not os.path.isdir(dir_path):
                    os.mkdir(dir_path)
                blank_para = doc.add_paragraph()
                run = blank_para.add_run()
                run.add_break()
                id = doc.add_paragraph().add_run(file_no)
                id.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                id = doc.add_paragraph().add_run(mr_no)
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                id = doc.add_paragraph().add_run(patient_name)
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                id = doc.add_paragraph().add_run(dob)
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                report_type = re.sub('.png', '.docx', str(qr_code))
                report_type = report_type.lower()
                doc_name = report_type + '.docx'
                doc.save(os.path.join(dir_path, doc_name))
                return dir_path

    @staticmethod
    def create_dummy_pdf_for_reports(report_type, dir_path):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', size=28)
        text_letter = re.sub('[^a-z_]', '', str(report_type))
        for i in range(5):
            text = text_letter[2:] + '_' + str(i)
            pdf.cell(200, 10, txt=text, ln=1, align='C')
            pdf.add_page()
            report_type = re.sub('[^a-z_]', '', str(report_type))
        pdf_name = report_type[2:] + '.pdf'
        pdf.output(os.path.join(dir_path, pdf_name))

    def add_dummy_reports_(self):
        reports = os.listdir(self.qr_code_id_folder_path)
        for report in reports:
            dir_name = re.sub('.docx', '', str(report))
            doc_path = os.path.join(self.qr_code_id_folder_path, report)
            # pdf_report_name = re.sub('[^0-9_]', '', str(report))
            pdf_report_name = dir_name + '_code.pdf'
            pdf_path = os.path.join(self.destination_path, pdf_report_name)
            convert(doc_path, pdf_path)
            self.create_dummy_pdf_for_reports(dir_name, self.destination_path)

