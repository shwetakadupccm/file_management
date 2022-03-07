import os
import re
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx2pdf import convert
from fpdf import FPDF

class DataDigitization:

    def __init__(self, report_type_df, qr_code_path, master_list, qr_code_id_folder_path, destination_path):
        self.report_type_df = report_type_df
        self.qr_code_path = qr_code_path
        self.master_list = master_list
        self.qr_code_id_folder_path = qr_code_id_folder_path
        self.destination_path = destination_path

    def add_qr_code_in_word_doc(self):
        for index, report_type in enumerate(self.report_type_df['report_types']):
            report_type_no = index + 1
            doc = Document()
            doc.add_picture(self.qr_code_path)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            type = report_type
            text = doc.add_paragraph()
            report_type_name = text.add_run(str(report_type_no) + '. ' + type)
            report_type_name.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            report_type_name.bold = True
            report_type_name.font.size = Pt(28)
            report_type_name.font.name = 'Arial Black'
            # doc.save(os.path.join(destination_path, report_type + '.docx'))
            for i in range(len(self.master_list)):
                file_no = 'File_number: ' + str(self.master_list['file_number'][i])
                mr_no = 'MR_number: ' + str(self.master_list['mr_number'][i])
                patient_name = 'Patient_name: ' + str(self.master_list['patient_name'][i])
                dob = 'Date_of_Birth: ' + str(self.master_list['date_of_birth'][i])
                # dir = self.master_list['file_number'][i]
                # dir = re.sub('/', '_', str(dir))
                # dir_path = os.path.join(self.qr_code_id_folder_path, dir)
                # if not os.path.isdir(dir_path):
                #     os.mkdir(dir_path)
                blank_para = doc.add_paragraph()
                run = blank_para.add_run()
                run.add_break()
                id = doc.add_paragraph().add_run(file_no)
                id.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                id = doc.add_paragraph().add_run(mr_no)
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                id = doc.add_paragraph().add_run(patient_name)
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                id = doc.add_paragraph().add_run(dob)
                id.font.size = Pt(20)
                id.font.name = 'Arial Black'
                report_type = re.sub(' ', '_', report_type)
                report_type = report_type.lower()
                doc_name = str(report_type_no) + '_' + report_type + '.docx'
                doc.save(os.path.join(self.qr_code_id_folder_path, doc_name))

    @staticmethod
    def create_dummy_images_for_reports(report_type, dir_path):
        for i in range(5):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('Arial', size=28)
            text_letter = re.sub('[^a-z_]', '', str(report_type))
            print(text_letter)
            text = text_letter[1:] + '_' + str(i)
            pdf.cell(200, 10, txt=text, ln=1, align='C')
            report_type = re.sub('[^a-z_]', '', str(report_type))
            pdf_name = report_type[1:] + '_' + str(i) + '.pdf'
            pdf.output(os.path.join(dir_path, pdf_name))

    def add_dummy_reports(self):
        reports = os.listdir(self.qr_code_id_folder_path)
        for report in reports:
            dir_name = re.sub('.docx', '', str(report))
            each_report_dir = os.path.join(self.destination_path, dir_name)
            os.mkdir(each_report_dir)
            doc_path = os.path.join(self.qr_code_id_folder_path, report)
            pdf_report_name = re.sub('[^0-9]', '', str(report))
            pdf_report_name = pdf_report_name + '_code.pdf'
            pdf_path = os.path.join(each_report_dir, pdf_report_name)
            convert(doc_path, pdf_path)
            self.create_dummy_images_for_reports(dir_name, each_report_dir)
